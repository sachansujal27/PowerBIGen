from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DOCXReportGenerator:

    def __init__(self, report, output_path):
        self.report = report
        self.output_path = output_path

    def add_heading(self, document, text, level=1):

        heading = document.add_heading("", level=level)

        run = heading.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 70, 140)

    def generate(self):

        document = Document()

        # -------------------------
        # Title
        # -------------------------

        title = document.add_heading("", level=0)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run("Business Analysis Report")

        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 70, 140)

        document.add_paragraph(
            "AI Generated Business Intelligence Report"
        )

        document.add_paragraph(
            f"Generated Date : {self.report.get('generated_date','')}"
        )

        document.add_page_break()

        # -------------------------
        # Executive Summary
        # -------------------------

        self.add_heading(document, "Executive Summary")

        document.add_paragraph(
            self.report.get(
                "executive_summary",
                "No executive summary available."
            )
        )

        # -------------------------
        # Dataset Information
        # -------------------------

        self.add_heading(document, "Dataset Information")

        info = self.report.get(
            "dataset_information",
            {}
        )

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        hdr = table.rows[0].cells

        hdr[0].text = "Property"
        hdr[1].text = "Value"

        for key, value in info.items():

            row = table.add_row().cells

            row[0].text = str(key)
            row[1].text = str(value)

        # -------------------------
        # KPI
        # -------------------------

        self.add_heading(document, "Key Performance Indicators")

        kpis = self.report.get(
            "kpis",
            []
        )

        if len(kpis):

            table = document.add_table(
                rows=1,
                cols=2
            )

            table.style = "Colorful Grid"

            hdr = table.rows[0].cells

            hdr[0].text = "KPI"
            hdr[1].text = "Value"

            for item in kpis:

                row = table.add_row().cells

                row[0].text = item.get(
                    "title",
                    ""
                )

                row[1].text = str(
                    item.get(
                        "value",
                        ""
                    )
                )

        document.add_page_break()

        # -------------------------
        # Business Findings
        # -------------------------

        self.add_heading(document, "Business Findings")

        insights = self.report.get(
            "insights",
            []
        )

        for item in insights:

            if isinstance(item, dict):

                p = document.add_paragraph(style="List Bullet")

                p.add_run(
                    f"{item.get('column')} : "
                ).bold = True

                p.add_run(
                    f"Total = {item.get('total')}, "
                    f"Average = {item.get('average')}, "
                    f"Maximum = {item.get('maximum')}, "
                    f"Minimum = {item.get('minimum')}"
                )

            else:

                document.add_paragraph(
                    str(item),
                    style="List Bullet"
                )

        # -------------------------
        # Risks
        # -------------------------

        self.add_heading(document, "Business Risks")

        for risk in self.report.get(
            "risks",
            []
        ):

            document.add_paragraph(
                risk,
                style="List Bullet"
            )

        # -------------------------
        # Opportunities
        # -------------------------

        self.add_heading(document, "Business Opportunities")

        for item in self.report.get(
            "opportunities",
            []
        ):

            document.add_paragraph(
                item,
                style="List Bullet"
            )

        # -------------------------
        # Recommendations
        # -------------------------

        self.add_heading(document, "Recommendations")

        for rec in self.report.get(
            "recommendations",
            []
        ):

            document.add_paragraph(
                rec,
                style="List Bullet"
            )

        # -------------------------
        # Conclusion
        # -------------------------

        self.add_heading(document, "Conclusion")

        document.add_paragraph(
            self.report.get(
                "conclusion",
                ""
            )
        )

        document.add_page_break()

        # -------------------------
        # Footer
        # -------------------------

        p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(
            "Generated by PowerBIGen"
        )

        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(
            120,
            120,
            120
        )

        document.save(
            self.output_path
        )